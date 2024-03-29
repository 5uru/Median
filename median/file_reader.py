import io

import pypdf
from docx import Document

from median.utils import median_logger


def read_docx(docx_file):
    """
    Read the content of a DOCX file.

    :param docx_file: The path to the DOCX file or a file-like object.
    :return: The content of the DOCX file as a string.
    """
    if isinstance(docx_file, str):
        doc = Document(docx_file)
    else:
        doc = Document(io.BytesIO(docx_file.read()))
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])


def read_pdf(pdf_file):
    """
    Read the content of a PDF file with improved error handling and efficiency.

    :param pdf_file: The path to the PDF file or a file-like object.
    :return: The content of the PDF file as a string.
    """
    try:
        if isinstance(pdf_file, str):
            with open(pdf_file, "rb") as f:
                pdf_reader = pypdf.PdfReader(f)
        else:
            pdf_reader = pypdf.PdfReader(io.BytesIO(pdf_file.read()))

        pages_text = [
            pdf_reader.pages[page].extract_text()
            for page in range(len(pdf_reader.pages))
        ]
        return "".join(pages_text)
    except Exception as e:  # Consider catching more specific exceptions
        median_logger.error(f"Error reading PDF file: {e}")
        return None


def main(file, file_type):
    """
    Read the content of a file based on its type.

    :param file: The file object or the file path.
    :param file_type: The MIME type of the file.
    :return: The content of the file as a string or None if the file type is not supported.
    """
    file_readers = {
        "text/markdown": lambda f: f.read().decode(),
        "application/pdf": read_pdf,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": read_docx,
        "text/plain": lambda f: f.read().decode(),
    }

    if file_type not in file_readers:
        raise ValueError(f"Unsupported file type: {file_type}")

    return file_readers[file_type](file)
